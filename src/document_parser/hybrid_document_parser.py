"""
混合文档解析器 - 统一格式输出
"""
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
from pathlib import Path
import json
import pdfplumber

from .base import DocumentParser, ParsedDocument, ParsedSection


@dataclass
class ContentBlock:
    """统一的内容块"""
    block_type: str  # "title" / "text" / "table"
    content: Any  # 内容
    metadata: Dict[str, Any] = None


class HybridDocumentParser(DocumentParser):
    """
    混合文档解析器

    策略：
    1. PDF纯文本 → pdfplumber
    2. PDF有表格 → pdfplumber.extract_tables()
    3. PDF有图片/扫描版 → OCR降级
    4. DOCX → python-docx
    5. 输出统一格式 blocks数组
    """

    def __init__(self, debug: bool = False):
        self.debug = debug
        self.ocr = None
        self.title_library = self._load_title_library()
        self._init_ocr()

    def _load_title_library(self) -> set:
        """加载标题库"""
        try:
            lib_path = Path(__file__).parent / 'title_library.json'
            with open(lib_path, 'r', encoding='utf-8') as f:
                titles = json.load(f)
                return set(titles)
        except Exception as e:
            if self.debug:
                print(f"⚠️ 加载标题库失败: {e}")
            return set()

    def _normalize_title(self, title: str) -> str:
        """规范化标题：去掉数字前缀（如 "1. " 或 "10. "）"""
        import re
        # 匹配开头的数字和点（如 "1. " 或 "10. "）
        normalized = re.sub(r'^\d+\.\s+', '', title.strip())
        return normalized

    def _is_title(self, text: str) -> bool:
        """检查文本是否是标题"""
        normalized = self._normalize_title(text)
        return normalized in self.title_library

    def _init_ocr(self):
        """初始化OCR"""
        try:
            from paddleocr import PaddleOCR
            self.ocr = PaddleOCR(use_textline_orientation=True, lang="en")
            if self.debug:
                print("[OCR] 初始化成功")
        except Exception as e:
            if self.debug:
                print(f"⚠️ OCR初始化失败: {e}")

    def parse(self, file_path: str, **kwargs) -> ParsedDocument:
        """解析文件"""
        file_path = str(file_path)
        ext = Path(file_path).suffix.lower()

        blocks = []

        if ext == '.pdf':
            blocks = self._parse_pdf(file_path)
        elif ext == '.docx':
            blocks = self._parse_docx(file_path)
        elif ext == '.pptx':
            blocks = self._parse_pptx(file_path)
        else:
            return self._unknown_result(file_path)

        # 转换成ParsedSection
        sections = self._blocks_to_sections(blocks)

        return ParsedDocument(
            file_name=Path(file_path).name,
            file_type=ext.lstrip('.'),
            sections=sections,
            raw_params={"blocks_count": len(blocks)}
        )

    def _detect_heuristic_headings(self, page) -> set:
        """
        Detect heading lines by font size: lines whose average word size >= 16
        are treated as headings (e.g. UKRI PDFs use Merriweather-18pt for section
        titles vs Roboto-14pt for body text).  Also catches bold-font lines.
        Returns a set of normalised line-text strings.
        """
        import re
        from collections import defaultdict
        try:
            words = page.extract_words(extra_attrs=["fontname", "size"])
            if not words:
                return set()

            # Group words into lines by rounding their top-y to the nearest 3pt
            lines: dict = defaultdict(list)
            for w in words:
                y_key = round(float(w.get("top", 0)) / 3) * 3
                lines[y_key].append(w)

            heading_lines: set = set()
            for _, line_words in sorted(lines.items()):
                sizes = [float(w.get("size", 0)) for w in line_words]
                avg_size = sum(sizes) / len(sizes) if sizes else 0
                is_large = avg_size >= 16
                is_bold  = any("bold" in w.get("fontname", "").lower()
                               for w in line_words)
                if not (is_large or is_bold):
                    continue

                raw_text = " ".join(w["text"] for w in line_words)
                # Normalise internal whitespace to match extract_text() output
                line_text = re.sub(r"\s+", " ", raw_text).strip()
                if not line_text or len(line_text) > 100:
                    continue

                is_numbered = bool(re.match(r"^\d+[\.\)]\s", line_text))
                word_count  = len(line_text.split())
                # Keep numbered headings OR multi-word large-font lines
                if is_numbered or word_count >= 2:
                    heading_lines.add(line_text)

            return heading_lines
        except Exception:
            return set()

    def _parse_pdf(self, file_path: str) -> List[ContentBlock]:
        """PDF混合解析"""
        blocks = []

        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    # 步骤1：尝试提取表格
                    tables = page.extract_tables()
                    if tables:
                        if self.debug:
                            print(f"  Page {page_num}: 检测到{len(tables)}个表格")
                        for table in tables:
                            blocks.append(ContentBlock(
                                block_type="table",
                                content=table,
                                metadata={"page": page_num, "method": "pdfplumber"}
                            ))

                    # 步骤2：提取文本，并用字体大小启发式检测标题
                    text = page.extract_text()
                    if text:
                        if self.debug:
                            print(f"  Page {page_num}: 提取纯文本")
                        extra_headings = self._detect_heuristic_headings(page)
                        page_blocks = self._process_pdf_text(text, page_num, extra_headings)
                        blocks.extend(page_blocks)
                    else:
                        # 纯文本提取失败，尝试OCR
                        if self.ocr:
                            if self.debug:
                                print(f"  Page {page_num}: 纯文本为空，使用OCR")
                            page_blocks = self._parse_pdf_page_with_ocr(page, page_num)
                            blocks.extend(page_blocks)

        except Exception as e:
            if self.debug:
                print(f"❌ PDF解析失败: {e}")

        return blocks

    def _process_pdf_text(self, text: str, page_num: int,
                          extra_headings: set = None) -> List[ContentBlock]:
        """
        处理PDF文本：识别标题和文本块。

        Fixes applied vs original:
          - Bug B: text accumulated before the *first* recognised heading is no
            longer silently dropped (``current_title`` may be None).
          - extra_headings: font-size-detected heading lines are also accepted
            as section boundaries even when not in title_library.
        """
        import re
        blocks = []
        lines = text.split('\n')

        current_title: str | None = None
        current_text: list = []

        for line in lines:
            line = line.strip()
            if not line:
                continue

            norm_line = re.sub(r"\s+", " ", line)
            is_heading = (
                self._is_title(norm_line)
                or (extra_headings and norm_line in extra_headings)
            )

            if is_heading:
                # Bug B fix: flush accumulated text even if no title matched yet
                if current_text:
                    blocks.append(ContentBlock(
                        block_type="text",
                        content="\n".join(current_text),
                        metadata={"page": page_num, "title": current_title or ""}
                    ))
                current_text = []

                blocks.append(ContentBlock(
                    block_type="title",
                    content=norm_line,
                    metadata={"page": page_num}
                ))
                current_title = norm_line
            else:
                current_text.append(line)

        # Bug B fix: flush last block even if current_title is None
        if current_text:
            blocks.append(ContentBlock(
                block_type="text",
                content="\n".join(current_text),
                metadata={"page": page_num, "title": current_title or ""}
            ))

        return blocks

    def _parse_pdf_page_with_ocr(self, page, page_num: int) -> List[ContentBlock]:
        """用OCR解析PDF页面"""
        blocks = []

        try:
            im = page.to_image()
            pil_image = im.original

            ocr_result = self.ocr.ocr(pil_image, cls=True)

            if ocr_result:
                blocks = self._process_ocr_result(ocr_result, page_num)

        except Exception as e:
            if self.debug:
                print(f"⚠️ OCR处理失败: {e}")

        return blocks

    def _process_ocr_result(self, ocr_result: List, page_num: int) -> List[ContentBlock]:
        """处理OCR结果"""
        blocks = []
        current_title = None
        current_text = []

        for line in ocr_result:
            if not line:
                continue

            line_text = " ".join([word[0] for word in line])

            # 检查是否是标题
            if self._is_title(line_text):
                if current_text and current_title:
                    blocks.append(ContentBlock(
                        block_type="text",
                        content="\n".join(current_text),
                        metadata={"page": page_num, "title": current_title, "method": "ocr"}
                    ))
                current_text = []

                blocks.append(ContentBlock(
                    block_type="title",
                    content=line_text,
                    metadata={"page": page_num, "method": "ocr"}
                ))
                current_title = line_text
            else:
                if line_text.strip():
                    current_text.append(line_text)

        if current_text:
            blocks.append(ContentBlock(
                block_type="text",
                content="\n".join(current_text),
                metadata={"page": page_num, "title": current_title or "", "method": "ocr"}
            ))

        return blocks

    def _parse_docx(self, file_path: str) -> List[ContentBlock]:
        """解析DOCX"""
        from docx import Document

        blocks = []

        try:
            doc = Document(file_path)

            current_title = None
            current_text = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                # 启发式识别标题：优先检查段落样式
                is_title = False

                # 1. 检查段落样式是否是Heading
                if para.style and para.style.name.startswith('Heading'):
                    is_title = True

                # 2. 再检查标题库
                elif self._is_title(text):
                    is_title = True

                if is_title:
                    # Bug B fix: flush accumulated text even if no title matched yet
                    if current_text:
                        blocks.append(ContentBlock(
                            block_type="text",
                            content="\n".join(current_text),
                            metadata={"title": current_title or ""}
                        ))
                    current_text = []

                    # 添加标题
                    blocks.append(ContentBlock(
                        block_type="title",
                        content=text,
                        metadata={}
                    ))
                    current_title = text
                else:
                    # 文本
                    current_text.append(text)

            # Bug B fix: flush last block even if current_title is None
            if current_text:
                blocks.append(ContentBlock(
                    block_type="text",
                    content="\n".join(current_text),
                    metadata={"title": current_title or ""}
                ))

            # 处理表格
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)

                blocks.append(ContentBlock(
                    block_type="table",
                    content=table_data,
                    metadata={"method": "docx"}
                ))

        except Exception as e:
            if self.debug:
                print(f"❌ DOCX解析失败: {e}")

        return blocks

    def _parse_pptx(self, file_path: str) -> List[ContentBlock]:
        """解析PPT"""
        # TODO: 实现PPT解析
        return []

    def _blocks_to_sections(self, blocks: List[ContentBlock]) -> List[ParsedSection]:
        """
        将ContentBlock转换为ParsedSection。

        Bug A fix: the original code set ``title = block.block_type``, which
        stored the literal string "title"/"text"/"table" instead of the actual
        heading text.  Now title-type blocks use their content as the title,
        and text/table blocks inherit the heading from their metadata.
        """
        sections = []

        for block in blocks:
            if block.block_type == "title":
                title = str(block.content) if block.content is not None else ""
            else:
                meta  = block.metadata or {}
                title = meta.get("title") or ""
                if not isinstance(title, str):
                    title = str(title)

            sections.append(ParsedSection(
                title=title,
                content=block.content,
                type=block.block_type
            ))

        return sections

    def auto_detect_params(self, file_path: str) -> Dict[str, Any]:
        """检测参数"""
        return {"method": "hybrid"}

    def _unknown_result(self, file_path: str) -> ParsedDocument:
        """未知文件类型"""
        return ParsedDocument(
            file_name=Path(file_path).name,
            file_type="unknown",
            sections=[],
            raw_params={"error": "不支持的文件类型"}
        )
