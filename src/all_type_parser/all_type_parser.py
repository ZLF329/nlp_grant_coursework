"""
All-Type Parser

Entry point for parsing any grant application file into the unified JSON format.

Pipeline for PDF files:
  1. fellowships_parser  — blue-box fellowship format (NIHR DF/AF)
     → if result is non-empty, done.
  2. RfPB_parser         — RfPB format (fast pre-check on page 1)
     → if result is non-empty, done.
  3. pdf_parser          — generic big-box PDF format
     → if result is non-empty, done.
  4. RfPB_parser         — second attempt as fallback
     → if result is non-empty, done.
  5. llm_fallback_parser — last resort: pdfplumber text + glm-ocr → qwen

Pipeline for DOCX files:
  1. python-docx raw text extraction
     → if >= MIN_CONTENT_CHARS, return as Raw Content.
  2. llm_fallback_parser — if text too sparse (< MIN_CONTENT_CHARS)

The output JSON always uses the same top-level keys as IC00458_after.json.
Only keys for which content was found are included.
"""
from __future__ import annotations

import json
import logging
import sys
from pathlib import Path
from typing import Optional

# Suppress noisy pdfminer FontBBox / encoding warnings (affects all parsers
# that use pdfplumber, including the _is_rfpb_pdf() pre-check below)
logging.getLogger("pdfminer").setLevel(logging.ERROR)

# ── path setup so sub-parsers can be imported regardless of cwd ──────────────
_SRC = Path(__file__).resolve().parent.parent   # …/src
if str(_SRC) not in sys.path:
    sys.path.insert(0, str(_SRC))


# ──────────────────────────── helpers ────────────────────────────────────────

def _is_empty(result: dict) -> bool:
    """Return True if the parser produced nothing useful."""
    if not result:
        return True
    # Ignore metadata-only keys (e.g. doc_type) when judging emptiness
    content = {k: v for k, v in result.items() if k != "doc_type"}
    if not content:
        return True
    # A non-empty dict must contain at least one non-empty value
    return all(
        (not v) or (isinstance(v, dict) and not v) or (isinstance(v, list) and not v)
        for v in content.values()
    )


def _total_text_length(result: dict) -> int:
    """Return total character count of all string leaf values in the result."""
    total = 0
    def _count(v):
        nonlocal total
        if isinstance(v, str):
            total += len(v)
        elif isinstance(v, dict):
            for child in v.values():
                _count(child)
        elif isinstance(v, list):
            for item in v:
                _count(item)
    _count(result)
    return total


# Minimum total characters for a DOCX result to be considered sufficient;
# below this threshold the LLM fallback is triggered.
_MIN_CONTENT_CHARS = 2000


def _json_output_path(input_path: str) -> str:
    """Derive the final JSON output path next to the input file."""
    p = Path(input_path)
    json_dir = p.parent / "json_data"
    json_dir.mkdir(exist_ok=True)
    return str(json_dir / (p.stem + ".json"))


def _save_json(data: dict, path: str) -> None:
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


# ──────────────────────────── stage 1 — fellowships_parser ───────────────────

def _try_fellowships_parser(pdf_path: str) -> dict:
    try:
        from .fellowships_parser import extract_all_sections
        result = extract_all_sections(pdf_path)
        return result if result else {}
    except Exception as e:
        print(f"[all_type_parser] fellowships_parser failed: {e}")
        return {}


# ──────────────────────────── stage 2 / 4 — RfPB_parser ─────────────────────

def _try_rfpb_parser(pdf_path: str) -> dict:
    try:
        from .RfPB_parser import extract_all_sections
        result = extract_all_sections(pdf_path)
        return result if result else {}
    except Exception as e:
        print(f"[all_type_parser] RfPB_parser failed: {e}")
        return {}


# ──────────────────────────── stage 3 — pdf_parser ───────────────────────────

def _try_pdf_parser(pdf_path: str) -> dict:
    try:
        from .pdf_parser import extract_all_big_box_sections
        result = extract_all_big_box_sections(pdf_path)
        return result if result else {}
    except Exception as e:
        print(f"[all_type_parser] pdf_parser failed: {e}")
        return {}


# ──────────────────────────── RfPB pre-check ─────────────────────────────────

def _is_rfpb_pdf(pdf_path: str, n_lines: int = 2) -> bool:
    """
    Return True if the first `n_lines` text lines of the PDF contain 'RfPB'
    (case-sensitive).  Uses pdfplumber on page 1 only — fast, no full parse.
    """
    try:
        import pdfplumber
        with pdfplumber.open(pdf_path) as pdf:
            if not pdf.pages:
                return False
            text = pdf.pages[0].extract_text() or ""
        first_lines = "\n".join(text.splitlines()[:n_lines])
        return "RfPB" in first_lines
    except Exception:
        return False


# ──────────────────────────── DOCX extraction ────────────────────────────────

def _try_docx_parse(docx_path: str) -> dict:
    """
    Extract all text from a DOCX file using python-docx and return a unified
    dict with the full content stored under APPLICATION DETAILS["Raw Content"].

    Returns {} if the file cannot be read or produces no text.
    """
    try:
        from docx import Document
        doc = Document(docx_path)
        parts: list = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    parts.append(row_text)

        full_text = "\n".join(parts)
        if not full_text.strip():
            return {}

        return {"APPLICATION DETAILS": {"Raw Content": full_text}}
    except Exception as e:
        print(f"[all_type_parser] docx parsing failed: {e}")
        return {}


# ──────────────────────────── stage 5 — llm_fallback_parser ─────────────────

def _try_llm_fallback(input_path: str) -> dict:
    """
    Last-resort parser: pdfplumber text → glm-ocr (image OCR) → qwen3.5:27b
    structured extraction via Ollama.  Also handles DOCX via python-docx.
    """
    try:
        from .llm_fallback_parser import extract_all_sections
        result = extract_all_sections(input_path)
        return result if result else {}
    except Exception as e:
        print(f"[all_type_parser] llm_fallback_parser failed: {e}")
        return {}


# ──────────────────────────── public API ─────────────────────────────────────

def parse(input_path: str) -> dict:
    """
    Parse any grant application file and return the unified JSON dict.

    For PDF files the four-stage rule-based pipeline is tried in order,
    with LLM as the final fallback.
    For DOCX files, python-docx extraction is used directly; LLM is
    triggered only if the extracted text is too sparse.
    """
    ext = Path(input_path).suffix.lower()

    if ext == ".pdf":
        # Fast pre-check: if page 1 mentions "RfPB", go straight to RfPB parser
        if _is_rfpb_pdf(input_path):
            print("[all_type_parser] detected RfPB PDF — using RfPB_parser directly")
            result = _try_rfpb_parser(input_path)
            if not _is_empty(result):
                print("[all_type_parser] ✓ RfPB_parser succeeded")
                return result
            print("[all_type_parser] RfPB_parser returned empty — falling back to LLM")
        else:
            # Stage 1: fellowship blue-box parser
            result = _try_fellowships_parser(input_path)
            if not _is_empty(result):
                print("[all_type_parser] ✓ fellowships_parser succeeded")
                return result

            # Stage 2: generic big-box PDF parser
            result = _try_pdf_parser(input_path)
            if not _is_empty(result):
                print("[all_type_parser] ✓ pdf_parser succeeded")
                return result

            # Stage 3: RfPB fallback for non-RfPB PDFs
            result = _try_rfpb_parser(input_path)
            if not _is_empty(result):
                print("[all_type_parser] ✓ RfPB_parser succeeded")
                return result

            print("[all_type_parser] all PDF parsers returned empty — falling back to LLM")

        # Stage 4: LLM fallback for all unrecognised PDFs
        print("[all_type_parser] falling back to LLM parser (glm-ocr + qwen3.5:27b)")
        result = _try_llm_fallback(input_path)
        if not _is_empty(result):
            print("[all_type_parser] ✓ llm_fallback_parser succeeded")
        else:
            print("[all_type_parser] ✗ all parsers returned empty")
        return result

    elif ext in (".docx", ".doc"):
        # DOCX: extract raw text with python-docx
        result = _try_docx_parse(input_path)
        if not _is_empty(result) and _total_text_length(result) >= _MIN_CONTENT_CHARS:
            print("[all_type_parser] ✓ docx parsing succeeded")
            return result

        if not _is_empty(result):
            print(
                f"[all_type_parser] docx result too sparse "
                f"({_total_text_length(result)} chars < {_MIN_CONTENT_CHARS}) "
                f"— falling back to LLM parser"
            )
        else:
            print("[all_type_parser] docx parsing returned empty — falling back to LLM parser")

        result = _try_llm_fallback(input_path)
        if not _is_empty(result):
            print("[all_type_parser] ✓ llm_fallback_parser succeeded")
        else:
            print("[all_type_parser] ✗ all parsers returned empty")
        return result

    else:
        # Unknown format — try LLM directly
        print(f"[all_type_parser] unsupported extension '{ext}' — trying LLM parser")
        result = _try_llm_fallback(input_path)
        if not _is_empty(result):
            print("[all_type_parser] ✓ llm_fallback_parser succeeded")
        else:
            print("[all_type_parser] ✗ all parsers returned empty")
        return result


def parse_and_save(input_path: str, output_path: Optional[str] = None) -> str:
    """
    Parse a file and write the unified JSON to disk.
    Returns the output JSON path.
    """
    result = parse(input_path)
    out = output_path or _json_output_path(input_path)
    _save_json(result, out)
    print(f"[all_type_parser] saved → {out}")
    return out


def parse_folder(folder_path: str, extensions: tuple = (".pdf", ".docx", ".doc")) -> list:
    """
    Parse all supported files in a folder and save each result to json_data/.
    Returns a list of output JSON paths.

    Parameters
    ----------
    folder_path : str
        Directory to scan (non-recursive).
    extensions : tuple
        File extensions to include. Defaults to pdf, docx, doc.
    """
    folder = Path(folder_path)
    files = [f for f in sorted(folder.iterdir()) if f.is_file() and f.suffix.lower() in extensions]

    if not files:
        print(f"[all_type_parser] no supported files found in {folder_path}")
        return []

    print(f"[all_type_parser] found {len(files)} file(s) in {folder_path}")
    saved: list = []
    for i, f in enumerate(files, 1):
        print(f"\n[all_type_parser] [{i}/{len(files)}] {f.name}")
        try:
            out = parse_and_save(str(f))
            saved.append(out)
        except Exception as e:
            print(f"[all_type_parser] ✗ failed: {e}")

    print(f"\n[all_type_parser] done — {len(saved)}/{len(files)} saved to json_data/")
    return saved


# ──────────────────────────── CLI ────────────────────────────────────────────

if __name__ == "__main__":
    import sys as _sys
    # Run as a module so relative imports work:
    #   python -m all_type_parser.all_type_parser <file>
    if len(_sys.argv) < 2:
        print("Usage: python -m all_type_parser.all_type_parser <input_file> [output.json]")
        _sys.exit(1)

    out = parse_and_save(_sys.argv[1], _sys.argv[2] if len(_sys.argv) > 2 else None)
    print(f"Done: {out}")
