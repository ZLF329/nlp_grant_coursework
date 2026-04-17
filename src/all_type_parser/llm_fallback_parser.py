"""
LLM Fallback Parser
===================
Last-resort parser for grant application PDFs that do not match any known
structural format (no blue-boxes, non-standard layout, scanned pages, etc.).

Strategy
--------
1. Try pdfplumber text extraction (fast, no GPU).
   - If text is rich enough (>= MIN_CHARS_PER_PAGE avg), send directly to LLM.
2. If text is sparse (scanned PDF), render pages to images via pdf2image and
   send them to glm-ocr (Ollama) for OCR, then concatenate the results.
3. Feed the extracted text to qwen3.5:27b (Ollama) and ask it to return a
   structured JSON matching the unified grant-application schema.

Environment variables (all optional — defaults match qwen3_ollama.py):
    OLLAMA_HOST          http://127.0.0.1:11434
    OLLAMA_MODEL         qwen3.5:27b      (structuring LLM)
    OLLAMA_OCR_MODEL     glm-ocr          (vision OCR model)
    OLLAMA_TIMEOUT       1200             (seconds)

Dependencies
------------
    pip install pdfplumber requests
    pip install pdf2image          # only needed for scanned PDFs
    # pdf2image also requires poppler on PATH:
    #   Ubuntu/Debian : apt install poppler-utils
    #   macOS         : brew install poppler
"""

from __future__ import annotations

import base64
import io
import json
import logging
import os
import re
from pathlib import Path
from typing import Optional

import pdfplumber
import requests

# Suppress noisy pdfminer FontBBox / encoding warnings
logging.getLogger("pdfminer").setLevel(logging.ERROR)

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

OLLAMA_HOST      = os.environ.get("OLLAMA_HOST",      "http://127.0.0.1:11434").rstrip("/")
OLLAMA_MODEL     = os.environ.get("OLLAMA_MODEL",     "qwen3.5:27b")
OLLAMA_OCR_MODEL = os.environ.get("OLLAMA_OCR_MODEL", "glm-ocr")
OLLAMA_TIMEOUT   = float(os.environ.get("OLLAMA_TIMEOUT", "1200"))

# Minimum average characters per page to skip OCR and use pdfplumber text directly.
# Below this threshold the page is treated as scanned and sent to glm-ocr.
MIN_CHARS_PER_PAGE: int = 300

# Page limits (token / cost guard)
MAX_PAGES_TEXT:  int = 60   # max pages whose text we send to the LLM

# DPI for pdf2image rendering
IMAGE_DPI: int = 150


# ---------------------------------------------------------------------------
# Stage 1 — text extraction via pdfplumber
# ---------------------------------------------------------------------------

def _extract_text_pdfplumber(pdf_path: str) -> tuple[str, float]:
    """
    Extract all text from a PDF using pdfplumber.

    Returns
    -------
    text : str
        Concatenated page texts (up to MAX_PAGES_TEXT pages).
    avg_chars_per_page : float
        Average character count per page — used to detect scanned PDFs.
    """
    pages_text: list[str] = []
    total_chars = 0

    with pdfplumber.open(pdf_path) as pdf:
        total_pages = len(pdf.pages)
        for page in pdf.pages[:MAX_PAGES_TEXT]:
            t = page.extract_text() or ""
            pages_text.append(t)
            total_chars += len(t)

    avg = total_chars / total_pages if total_pages else 0.0
    return "\n\n".join(pages_text), avg


def _extract_text_docx(docx_path: str) -> str:
    """
    Extract all text from a DOCX file using python-docx.

    Paragraphs and table cells are concatenated in document order.
    Returns a single string ready to be sent to the LLM.
    """
    from docx import Document

    doc = Document(docx_path)
    parts: list[str] = []

    # Paragraphs in document order
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Tables (append after paragraphs to preserve rough document order)
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)

    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Stage 2 — image OCR via glm-ocr (Ollama)
# ---------------------------------------------------------------------------

def _pdf_pages_to_base64(pdf_path: str) -> list[str]:
    """
    Render all PDF pages to PNG images and return them as base64-encoded strings.
    Requires pdf2image and poppler on PATH.
    """
    try:
        from pdf2image import convert_from_path
    except ImportError as exc:
        raise ImportError(
            "pdf2image is required for scanned PDF support. "
            "Install it with:  pip install pdf2image"
        ) from exc

    images = convert_from_path(pdf_path, dpi=IMAGE_DPI)
    result: list[str] = []
    for img in images:
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        result.append(base64.b64encode(buf.getvalue()).decode("utf-8"))
    return result


def _ocr_image_ollama(b64_image: str, page_num: int) -> str:
    """
    Send a single base64-encoded image to glm-ocr via Ollama and return the
    raw extracted text.
    """
    payload = {
        "model": OLLAMA_OCR_MODEL,
        "messages": [
            {
                "role": "user",
                "content": (
                    "Please extract all text from this image exactly as it appears, "
                    "preserving paragraph breaks and layout structure. "
                    "Output plain text only — no commentary."
                ),
                "images": [b64_image],
            }
        ],
        "stream": False,
        "options": {"temperature": 0.0},
    }
    try:
        resp = requests.post(
            f"{OLLAMA_HOST}/api/chat",
            json=payload,
            timeout=OLLAMA_TIMEOUT,
        )
        resp.raise_for_status()
        body = resp.json()
        return (body.get("message") or {}).get("content") or ""
    except Exception as e:
        print(f"[llm_fallback_parser] OCR failed for page {page_num}: {e}")
        return ""


def _ocr_pdf(pdf_path: str) -> str:
    """
    Convert PDF pages to images, OCR each one with glm-ocr, and concatenate
    the results into a single text string.
    """
    print("[llm_fallback_parser] sparse text detected — running image OCR (all pages)")
    b64_images = _pdf_pages_to_base64(pdf_path)
    page_texts: list[str] = []
    for i, b64 in enumerate(b64_images, 1):
        print(f"[llm_fallback_parser]   OCR page {i}/{len(b64_images)} …")
        text = _ocr_image_ollama(b64, page_num=i)
        if text.strip():
            page_texts.append(text)
    return "\n\n".join(page_texts)


# ---------------------------------------------------------------------------
# Stage 3 — structured extraction via qwen3.5:27b
# ---------------------------------------------------------------------------

# JSON schema for Ollama constrained generation
_UNIFIED_SCHEMA: dict = {
    "type": "object",
    "additionalProperties": False,
    "properties": {
        "doc_type": {"type": "string"},
        "SUMMARY INFORMATION": {
            "type": "object",
            "additionalProperties": False,
            "properties": {
                "Application Title": {"type": "string"},
                "Contracting Organisation": {"type": "string"},
                "Start Date": {"type": "string"},
                "End Date": {"type": "string"},
                "Duration (months)": {"type": "string"},
                "Total Cost to NIHR": {"type": "string"},
            },
            "required": [
                "Application Title",
                "Contracting Organisation",
                "Start Date",
                "End Date",
                "Duration (months)",
                "Total Cost to NIHR",
            ],
        },
        "LEAD APPLICANT & RESEARCH TEAM": {
            "type": "object",
            "additionalProperties": False,
            "properties": {
                "Lead Applicant": {
                    "type": ["object", "null"],
                    "additionalProperties": False,
                    "properties": {
                        "Full Name": {"type": "string"},
                        "Organisation": {"type": "string"},
                        "Department": {"type": "string"},
                        "Proposed Role": {"type": "string"},
                        "ORCID": {"type": "string"},
                        "% FTE Commitment": {"type": "string"},
                    },
                    "required": [
                        "Full Name",
                        "Organisation",
                        "Department",
                        "Proposed Role",
                        "ORCID",
                        "% FTE Commitment",
                    ],
                },
                "Joint Lead Applicant": {"type": ["object", "null"]},
                "Co-Applicants": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "Full Name": {"type": "string"},
                            "Proposed Role": {"type": "string"},
                            "Organisation": {"type": "string"},
                            "Department": {"type": "string"},
                            "ORCID": {"type": "string"},
                            "Position": {"type": "string"},
                        },
                        "required": [
                            "Full Name",
                            "Proposed Role",
                            "Organisation",
                            "Department",
                            "ORCID",
                            "Position",
                        ],
                    },
                },
            },
            "required": [
                "Lead Applicant",
                "Joint Lead Applicant",
                "Co-Applicants",
            ],
        },
        "APPLICATION DETAILS": {
            "type": "object",
            "additionalProperties": False,
            "properties": {
                "Plain English Summary of Research": {"type": "string"},
                "Plain English Summary": {"type": "string"},
                "Scientific Abstract": {"type": "string"},
                "Applicant CV": {"type": "string"},
                "Applicant Research Background": {"type": "string"},
                "Detailed Research Plan": {"type": "string"},
                "Changes from Previous Stage": {"type": "string"},
                "Patient & Public Involvement": {"type": "string"},
                "Working with People and Communities Summary": {"type": "string"},
                "Training & Development and Research Support": {"type": "string"},
                "SUPPORT AND MENTORSHIP": {"type": "string"},
                "Other Content": {"type": "string"},
            },
            "required": [
                "Plain English Summary of Research",
                "Plain English Summary",
                "Scientific Abstract",
                "Applicant CV",
                "Applicant Research Background",
                "Detailed Research Plan",
                "Changes from Previous Stage",
                "Patient & Public Involvement",
                "Working with People and Communities Summary",
                "Training & Development and Research Support",
                "SUPPORT AND MENTORSHIP",
                "Other Content",
            ],
        },
        "SUMMARY BUDGET": {"type": "string"},
    },
    "required": [
        "doc_type",
        "SUMMARY INFORMATION",
        "LEAD APPLICANT & RESEARCH TEAM",
        "APPLICATION DETAILS",
        "SUMMARY BUDGET",
    ],
}


_SYSTEM_PROMPT = """\
You are an expert at extracting structured information from grant application documents.
Given raw text from a grant application PDF, extract the information and return it as
valid JSON that matches the provided schema.

Rules:
- Return one JSON object that matches the schema exactly.
- Use the schema's field names exactly as written. Do not invent new keys.
- Always include every key required by the schema.
- If content is missing, use empty string "" for missing text fields, null for
  missing object fields, and [] for missing arrays.
- Do NOT invent or hallucinate any information.
- For narrative text fields, copy the original document text as literally as possible.
- Do NOT summarise, paraphrase, compress, rewrite, or simplify the original wording.
- Preserve original wording, ordering, and paragraph breaks whenever possible.
- Only make minimal fixes for obvious OCR noise such as broken spacing or duplicated punctuation.
- Map document section titles and near-synonyms into the predefined schema keys.
- Canonical APPLICATION DETAILS mappings:
  - lay summary headings -> "Plain English Summary of Research" or "Plain English Summary"
  - abstract headings -> "Scientific Abstract"
  - research plan / methods / proposal body headings -> "Detailed Research Plan"
  - first-stage revision headings -> "Changes from Previous Stage"
  - PPI headings -> "Patient & Public Involvement"
  - working with people / communities headings -> "Working with People and Communities Summary"
  - training / development / host support headings -> "Training & Development and Research Support"
  - supervisor / mentorship / support headings -> "SUPPORT AND MENTORSHIP"
  - any section that does not fit the above categories -> "Other Content" (concatenate all such sections here)
- Canonical SUMMARY INFORMATION mappings:
  - project or research title -> "Application Title"
  - host or contracting organisation -> "Contracting Organisation"
  - duration or grant duration in months -> "Duration (months)"
  - total cost / research costs -> "Total Cost to NIHR"
- Return ONLY the JSON object with no preamble, markdown fences, or explanation.
"""

_USER_PROMPT_TEMPLATE = """\
Extract structured grant application information from the text below and return it as JSON.

--- BEGIN DOCUMENT ---
{text}
--- END DOCUMENT ---
"""


def _strip_think_tags(text: str) -> str:
    """Remove <think>…</think> blocks emitted by reasoning models."""
    return re.sub(r"<think>.*?</think>\s*", "", text or "", flags=re.DOTALL).strip()


def _extract_json_object(text: str) -> str:
    """Pull the outermost { … } JSON object from a string."""
    clean = (text or "").strip()
    first = clean.find("{")
    last  = clean.rfind("}")
    if first != -1 and last != -1 and last > first:
        return clean[first : last + 1]
    return clean


def _structure_with_llm(text: str) -> dict:
    """
    Send extracted text to qwen3.5:27b via Ollama and parse the structured
    JSON response into a Python dict.

    Uses Ollama's `format` parameter for constrained JSON generation,
    matching the pattern in qwen3_ollama.py.
    """
    # Guard against extremely long documents exceeding context window
    user_text = text[:120_000]

    # Output budget: roughly 1 output token per 8 input chars, clamped to
    # [8192, 32768].  This avoids both truncation (the 8192-fixed bug) and
    # unnecessarily large allocations for short documents.
    num_predict = max(8192, min(len(user_text) // 8, 32768))

    payload = {
        "model": OLLAMA_MODEL,
        "messages": [
            {"role": "system", "content": _SYSTEM_PROMPT},
            {"role": "user",   "content": _USER_PROMPT_TEMPLATE.format(text=user_text)},
        ],
        "stream": False,
        "format": _UNIFIED_SCHEMA,
        "options": {
            "temperature": 0.1,
            "top_p": 0.9,
            "num_predict": num_predict,
        },
        "think": False,
    }

    try:
        resp = requests.post(
            f"{OLLAMA_HOST}/api/chat",
            json=payload,
            timeout=OLLAMA_TIMEOUT,
        )
        resp.raise_for_status()
    except requests.exceptions.ConnectionError as exc:
        raise RuntimeError(
            f"Could not connect to Ollama at {OLLAMA_HOST}. "
            "Make sure the server is running or set OLLAMA_HOST correctly."
        ) from exc
    except requests.exceptions.Timeout as exc:
        raise RuntimeError(
            f"Timed out waiting for Ollama at {OLLAMA_HOST}. "
            "Try increasing OLLAMA_TIMEOUT."
        ) from exc

    body = resp.json()
    raw = (body.get("message") or {}).get("content") or ""
    raw = _strip_think_tags(raw)
    raw = _extract_json_object(raw)

    try:
        return json.loads(raw)
    except json.JSONDecodeError as e:
        print(f"[llm_fallback_parser] JSON parse error: {e}")
        print(f"[llm_fallback_parser] raw LLM output length: {len(raw)} and the num_predict is: {num_predict}")
        print(f"[llm_fallback_parser] raw LLM output (first 500 chars): {raw[:500]}")
        return {}


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def extract_all_sections(input_path: str) -> dict:
    """
    Extract all grant-application sections from a PDF or DOCX file and return
    a unified JSON dict matching the IC00458_after.json schema.

    Pipeline for PDF
    ----------------
    1. pdfplumber text extraction
       - Rich text  (avg >= MIN_CHARS_PER_PAGE) → feed directly to LLM
       - Sparse text (scanned PDF)              → glm-ocr → feed to LLM
    2. qwen3.5:27b structured extraction → unified dict

    Pipeline for DOCX (and other non-PDF files)
    --------------------------------------------
    1. python-docx text extraction (paragraphs + tables)
    2. qwen3.5:27b structured extraction → unified dict
       (image OCR step is skipped — DOCX has selectable text by definition)

    Parameters
    ----------
    input_path : str
        Absolute or relative path to the input file (PDF or DOCX).

    Returns
    -------
    dict
        Unified JSON dict.  Empty dict if extraction failed completely.
    """
    ext = Path(input_path).suffix.lower()
    print(f"[llm_fallback_parser] processing: {input_path}")

    # ── DOCX branch: extract text with python-docx, skip image OCR ───────────
    if ext in (".docx", ".doc"):
        try:
            text = _extract_text_docx(input_path)
        except Exception as e:
            print(f"[llm_fallback_parser] python-docx extraction failed: {e}")
            return {}

        if not text.strip():
            print("[llm_fallback_parser] no text extracted from DOCX — giving up")
            return {}

        print(f"[llm_fallback_parser] DOCX: extracted {len(text):,} chars")

    # ── PDF branch: pdfplumber → (glm-ocr if sparse) ─────────────────────────
    else:
        # Stage 1: pdfplumber text extraction
        try:
            text, avg_chars = _extract_text_pdfplumber(input_path)
        except Exception as e:
            print(f"[llm_fallback_parser] pdfplumber failed: {e}")
            text, avg_chars = "", 0.0

        print(f"[llm_fallback_parser] avg chars/page: {avg_chars:.0f}")

        # Stage 2: image OCR if text is too sparse
        if avg_chars < MIN_CHARS_PER_PAGE:
            try:
                ocr_text = _ocr_pdf(input_path)
                if ocr_text.strip():
                    text = ocr_text
            except Exception as e:
                print(f"[llm_fallback_parser] image OCR failed: {e}")
                # Fall through — use whatever pdfplumber gave us (may be empty)

        if not text.strip():
            print("[llm_fallback_parser] no text available — giving up")
            return {}

    # ── Stage 3: structured extraction via LLM (both branches) ───────────────
    print(f"[llm_fallback_parser] sending {len(text):,} chars to {OLLAMA_MODEL} …")
    try:
        result = _structure_with_llm(text)
    except Exception as e:
        print(f"[llm_fallback_parser] LLM structuring failed: {e}")
        return {}

    if result:
        result.setdefault("doc_type", "llm_fallback")
        print("[llm_fallback_parser] ✓ extraction complete")
    else:
        print("[llm_fallback_parser] ✗ LLM returned empty result")

    return result
